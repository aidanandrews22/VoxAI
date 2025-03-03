"""
Document processor module for extracting text from PDF and Word documents.
"""
import io
import os
import tempfile
from typing import Dict, Any, List, Tuple

import fitz  # PyMuPDF
import docx
from PIL import Image

from app.core.config import settings
from app.core.logging import logger
from app.services.file_processors import FileProcessor


class DocumentProcessor(FileProcessor):
    """
    Base processor for document files.
    """
    
    async def process(self, file_content: bytes, file_path: str) -> str:
        """
        Process document content and extract text.
        
        Args:
            file_content: Raw document file content bytes
            file_path: Path to the document file
            
        Returns:
            str: Extracted text content
        """
        raise NotImplementedError("Subclasses must implement this method")
    
    async def get_metadata(self, file_content: bytes, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from document file.
        
        Args:
            file_content: Raw document file content bytes
            file_path: Path to the document file
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        return {
            'file_size': len(file_content),
            'file_extension': os.path.splitext(file_path)[1],
        }
    
    async def _extract_text_and_description_from_image(self, image: Image.Image) -> str:
        """
        Extract text and generate description from an image using Google Gemini.
        
        Args:
            image: PIL Image object
            
        Returns:
            str: Extracted text and generated description
        """
        try:
            from google import genai
            import os
            
            # Initialize the Gemini client
            client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
            
            # Convert PIL Image to bytes for the API
            with io.BytesIO() as output:
                image.save(output, format="PNG")
                image_bytes = output.getvalue()
            
            # Create the request parts
            from google.genai import types
            image_part = types.Part.from_bytes(data=image_bytes, mime_type="image/png")
            
            combined_prompt = """
                1. Extract all text visible in this image. If no text is visible, respond with 'No text detected.'
                2. Provide a detailed description of what's in this image.
            """

            response = client.models.generate_content(
                model="gemini-1.5-flash-8b",
                contents=[combined_prompt, image_part]
            )
            
            return response.text
        except Exception as e:
            logger.error(f"Error extracting text and description from image: {e}")
            return "Error processing image content." 


class PDFProcessor(DocumentProcessor):
    """
    Processor for PDF files.
    Handles: application/pdf
    """
    
    async def process(self, file_content: bytes, file_path: str) -> str:
        """
        Process PDF content and extract text.
        
        Args:
            file_content: Raw PDF file content bytes
            file_path: Path to the PDF file
            
        Returns:
            str: Extracted text content
        """
        try:
            # Process PDF using PyMuPDF (fitz)
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            
            full_text = []
            
            # Process all pages
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Extract text from the page
                page_text = page.get_text()
                
                # Extract images and their positions from the page
                image_info = await self._extract_images_with_positions(page)
                
                if not image_info and not page_text.strip():
                    # Empty page
                    full_text.append(f"[PAGE {page_num + 1} - EMPTY]")
                    continue
                
                if not image_info:
                    # Text only page
                    full_text.append(f"[PAGE {page_num + 1}]\n{page_text}")
                    continue
                
                if not page_text.strip():
                    # Image only page
                    image_texts = []
                    for img, _, _ in image_info:
                        img_text = await self._extract_text_and_description_from_image(img)
                        if img_text:
                            image_texts.append(f"[IMAGE CONTENT START]\n{img_text}\n[IMAGE CONTENT END]")
                    
                    full_text.append(f"[PAGE {page_num + 1} - IMAGES ONLY]\n\n" + "\n\n".join(image_texts))
                    continue
                
                # Page with both text and images - attempt to maintain order
                # For simplicity, we'll divide the page into top and bottom sections
                # and place images accordingly
                
                # Get page height
                page_height = page.rect.height
                
                # Separate images into top half and bottom half
                top_images = []
                bottom_images = []
                
                for img, x, y in image_info:
                    if y < page_height / 2:
                        top_images.append(img)
                    else:
                        bottom_images.append(img)
                
                # Process top images
                top_image_texts = []
                for img in top_images:
                    img_text = await self._extract_text_and_description_from_image(img)
                    if img_text:
                        top_image_texts.append(f"[IMAGE CONTENT START]\n{img_text}\n[IMAGE CONTENT END]")
                
                # Process bottom images
                bottom_image_texts = []
                for img in bottom_images:
                    img_text = await self._extract_text_and_description_from_image(img)
                    if img_text:
                        bottom_image_texts.append(f"[IMAGE CONTENT START]\n{img_text}\n[IMAGE CONTENT END]")
                
                # Combine content in a way that approximates the original layout
                page_content = []
                page_content.append(f"[PAGE {page_num + 1}]")
                
                if top_image_texts:
                    page_content.append("\n\n".join(top_image_texts))
                
                page_content.append(page_text)
                
                if bottom_image_texts:
                    page_content.append("\n\n".join(bottom_image_texts))
                
                full_text.append("\n\n".join(page_content))
            
            pdf_document.close()
            
            return "\n\n".join(full_text)
        except Exception as e:
            logger.error(f"Error processing PDF file: {e}")
            raise
    
    async def get_metadata(self, file_content: bytes, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from PDF file.
        
        Args:
            file_content: Raw PDF file content bytes
            file_path: Path to the PDF file
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        base_metadata = await super().get_metadata(file_content, file_path)
        
        try:
            # Process PDF using PyMuPDF (fitz)
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            
            # Count total images in the document
            total_images = 0
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                total_images += len(page.get_images(full=True))
            
            # Extract document metadata
            metadata = {
                'page_count': len(pdf_document),
                'image_count': total_images,
                'title': pdf_document.metadata.get('title', ''),
                'author': pdf_document.metadata.get('author', ''),
                'subject': pdf_document.metadata.get('subject', ''),
                'keywords': pdf_document.metadata.get('keywords', ''),
                'creator': pdf_document.metadata.get('creator', ''),
                'producer': pdf_document.metadata.get('producer', ''),
                'creation_date': pdf_document.metadata.get('creationDate', ''),
                'modification_date': pdf_document.metadata.get('modDate', ''),
            }
            
            # Remove empty metadata fields
            metadata = {k: v for k, v in metadata.items() if v}
            
            # Add to base metadata
            base_metadata.update(metadata)
            
            pdf_document.close()
            
            return base_metadata
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return base_metadata
    
    async def _extract_images_with_positions(self, page: fitz.Page) -> List[Tuple[Image.Image, float, float]]:
        """
        Extract images from a PDF page along with their positions.
        
        Args:
            page: PyMuPDF page object
            
        Returns:
            List[Tuple[Image.Image, float, float]]: List of tuples containing (image, x_pos, y_pos)
        """
        images_with_pos = []
        
        # Get image list
        img_list = page.get_images(full=True)
        
        for img_index, img_info in enumerate(img_list):
            try:
                xref = img_info[0]
                base_image = page.parent.extract_image(xref)
                
                # Convert to PIL Image
                image_data = base_image["image"]
                img = Image.open(io.BytesIO(image_data))
                
                # Only process images that are big enough to contain meaningful text
                if img.width > 100 and img.height > 100:
                    # Find the image on the page to get its position
                    for img_rect in page.get_image_rects(xref):
                        # Use the top-left corner as the position reference
                        x_pos = img_rect.x0
                        y_pos = img_rect.y0
                        images_with_pos.append((img, x_pos, y_pos))
                        break  # Just use the first occurrence if multiple
            except Exception as e:
                logger.error(f"Error extracting image {img_index} from PDF: {e}")
        
        # Sort by y-position (top to bottom)
        images_with_pos.sort(key=lambda x: x[2])
        
        return images_with_pos
    
    def _extract_images_from_page(self, page: fitz.Page) -> List[Image.Image]:
        """
        Extract images from a PDF page.
        
        Args:
            page: PyMuPDF page object
            
        Returns:
            List[Image.Image]: List of extracted images as PIL Image objects
        """
        images = []
        
        # Get image list
        img_list = page.get_images(full=True)
        
        for img_index, img_info in enumerate(img_list):
            try:
                xref = img_info[0]
                base_image = page.parent.extract_image(xref)
                
                # Convert to PIL Image
                image_data = base_image["image"]
                img = Image.open(io.BytesIO(image_data))
                
                # Only process images that are big enough to contain meaningful text
                if img.width > 100 and img.height > 100:
                    images.append(img)
            except Exception as e:
                logger.error(f"Error extracting image {img_index} from PDF: {e}")
        
        return images


class WordProcessor(DocumentProcessor):
    """
    Processor for Word documents.
    Handles: application/vnd.openxmlformats-officedocument.wordprocessingml.document
    """
    
    async def process(self, file_content: bytes, file_path: str) -> str:
        """
        Process Word document content and extract text.
        
        Args:
            file_content: Raw Word document file content bytes
            file_path: Path to the Word document file
            
        Returns:
            str: Extracted text content
        """
        try:
            # Write content to a temporary file because python-docx requires a file path
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                temp_path = temp_file.name
            
            # Process the Word document
            doc = docx.Document(temp_path)
            
            # Extract text from paragraphs and tables
            document_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    document_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        table_content.append(" | ".join(row_text))
                if table_content:
                    document_content.append("\n".join(table_content))
            
            # Clean up the temporary file
            os.unlink(temp_path)
            
            # Process images if there are any
            images = await self._extract_images_from_doc(file_content)
            for img in images:
                img_text = await self._extract_text_and_description_from_image(img)
                if img_text:
                    # Add a marker to indicate this is from an image
                    document_content.append(f"[IMAGE CONTENT START]\n{img_text}\n[IMAGE CONTENT END]")
            
            return "\n\n".join(document_content)
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            raise
    
    async def get_metadata(self, file_content: bytes, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from Word document.
        
        Args:
            file_content: Raw Word document file content bytes
            file_path: Path to the Word document file
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        base_metadata = await super().get_metadata(file_content, file_path)
        
        try:
            # Write content to a temporary file because python-docx requires a file path
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                temp_path = temp_file.name
            
            # Process the Word document
            doc = docx.Document(temp_path)
            
            # Count images in the document
            images = await self._extract_images_from_doc(file_content)
            image_count = len(images)
            
            # Extract document properties
            core_properties = doc.core_properties
            
            metadata = {
                'title': core_properties.title,
                'author': core_properties.author,
                'subject': core_properties.subject,
                'keywords': core_properties.keywords,
                'created': core_properties.created.isoformat() if core_properties.created else None,
                'modified': core_properties.modified.isoformat() if core_properties.modified else None,
                'last_modified_by': core_properties.last_modified_by,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'image_count': image_count,
            }
            
            # Remove empty metadata fields
            metadata = {k: v for k, v in metadata.items() if v}
            
            # Add to base metadata
            base_metadata.update(metadata)
            
            # Clean up the temporary file
            os.unlink(temp_path)
            
            return base_metadata
        except Exception as e:
            logger.error(f"Error extracting Word document metadata: {e}")
            return base_metadata
    
    async def _extract_images_from_doc(self, file_content: bytes) -> List[Image.Image]:
        """
        Extract images from a Word document.
        
        Args:
            file_content: Raw Word document file content bytes
            
        Returns:
            List[Image.Image]: List of extracted images as PIL Image objects
        """
        images = []
        
        try:
            # Write content to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                temp_path = temp_file.name
            
            # Use zipfile to extract media from docx (which is a zip file)
            import zipfile
            
            with zipfile.ZipFile(temp_path) as doc_zip:
                # Find all image files in the zip
                image_files = [f for f in doc_zip.namelist() if f.startswith('word/media/')]
                
                for image_path in image_files:
                    try:
                        image_data = doc_zip.read(image_path)
                        img = Image.open(io.BytesIO(image_data))
                        # Only process images that are big enough to contain meaningful text
                        if img.width > 100 and img.height > 100:
                            images.append(img)
                    except Exception as e:
                        logger.error(f"Error extracting image {image_path} from Word document: {e}")
            
            # Clean up
            os.unlink(temp_path)
            
            return images
        except Exception as e:
            logger.error(f"Error extracting images from Word document: {e}")
            return images 